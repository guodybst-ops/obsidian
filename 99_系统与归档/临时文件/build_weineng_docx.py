from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\89836\Documents\Obsidian Vault\05_内容输出\内容\公众号草稿\杭州蔚能海洋宣传稿-编辑包-2026-07-15.md")
OUTPUT = Path(r"C:\Users\89836\Documents\Obsidian Vault\05_内容输出\内容\公众号草稿\杭州蔚能海洋宣传稿-编辑包-2026-07-15.docx")

FONT = "Microsoft YaHei"
INK = RGBColor(33, 43, 54)
BLUE = RGBColor(25, 92, 128)
DEEP_BLUE = RGBColor(14, 62, 91)
MUTED = RGBColor(100, 112, 122)
LIGHT_BLUE = "E8F1F5"
PALE_BLUE = "F4F8FA"
LINE = "C9D8E0"
WARN = RGBColor(157, 103, 20)


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_fixed_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, 8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.38

    for name, size, color, before, after in (
        ("Heading 1", 16, DEEP_BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DEEP_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption Text" not in doc.styles:
        caption = doc.styles.add_style("Caption Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Caption Text"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(9)
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("杭州蔚能海洋科技有限公司  |  宣传稿编辑包")
    set_font(r, 8.5, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])


def add_rich_text(paragraph, text, size=None, color=INK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, (size or 10.5) - 0.5, color=MUTED)
        else:
            run = paragraph.add_run(part)
            set_font(run, size, color=color)


def add_numbered_item(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.20)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.30
    add_rich_text(p, text)


def add_bullet_item(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.20)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.28
    add_rich_text(p, text)


def add_note_box(doc, text, label="核稿提示"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, 120, 160, 120, 160)
    set_table_borders(table, color=LINE, size=5)
    set_fixed_table_geometry(table, [9360], indent_dxa=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}：")
    set_font(run, 10, bold=True, color=BLUE)
    add_rich_text(p, text, 10, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_record_blocks(doc, headers, rows):
    for idx, row in enumerate(rows, 1):
        title = row[1] if len(row) > 1 else f"条目 {idx}"
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.space_before = Pt(9 if idx > 1 else 3)
        p.add_run(f"{idx:02d}  {title}")

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for col_idx, value in enumerate(row):
            if col_idx == 1 or not value:
                continue
            cells = table.add_row().cells
            label = headers[col_idx] if col_idx < len(headers) else f"字段{col_idx + 1}"
            cells[0].text = label
            cells[1].text = value
            set_cell_shading(cells[0], LIGHT_BLUE)
            for cell in cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, 85, 120, 85, 120)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.18
                    for run in paragraph.runs:
                        set_font(run, 9.2, bold=(cell is cells[0]), color=(BLUE if cell is cells[0] else INK))
        set_table_borders(table, color=LINE, size=5)
        set_fixed_table_geometry(table, [1700, 7660], indent_dxa=120)


def add_title_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(34)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("杭州蔚能海洋科技有限公司")
    set_font(r, 16, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("宣传稿编辑包")
    set_font(r, 28, bold=True, color=DEEP_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("杭州良渚新城微信公众号 · 送企业核稿版")
    set_font(r, 11.5, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata = [
        ("稿件状态", "待企业核稿"),
        ("发布视角", "良渚新城管委会"),
        ("成稿日期", "2026年7月15日"),
        ("材料依据", "用户提供的企业资料及公开信息线索"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 110, 140, 110, 140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                set_font(run, 10, bold=(idx == 0), color=(BLUE if idx == 0 else INK))
    set_table_borders(table, color=LINE, size=5)
    set_fixed_table_geometry(table, [2100, 6000], indent_dxa=630)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    add_note_box(
        doc,
        "本稿依据现有材料形成。SIREN-38列装、企业成立时间、良渚落地情况、合作资质及图片授权等事项，须在发布前完成企业书面确认。",
        "使用说明",
    )
    doc.add_page_break()


def build():
    text = SOURCE.read_text(encoding="utf-8")
    text = re.sub(r"\A---\n.*?\n---\n", "", text, flags=re.S)
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    i = 0
    first_source_heading_skipped = False
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\-:| ]+\|$", lines[i + 1].strip()):
            headers = split_table_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_record_blocks(doc, headers, rows)
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            if not first_source_heading_skipped and title == "杭州蔚能海洋科技有限公司宣传稿编辑包":
                first_source_heading_skipped = True
                i += 1
                continue
            p = doc.add_paragraph(style="Heading 1")
            add_rich_text(p, title, 16, DEEP_BLUE)
        elif line.startswith("## "):
            title = line[3:].strip()
            if title == "四、公众号正文":
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_rich_text(p, title, 16, DEEP_BLUE)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_rich_text(p, line[4:].strip(), 13, BLUE)
        elif re.match(r"^\d+\. ", line):
            add_numbered_item(doc, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            add_bullet_item(doc, line[2:].strip())
        elif line.startswith("[配图"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line.strip("[]"))
            set_font(r, 9.3, italic=True, color=WARN)
        elif line.startswith("▲"):
            p = doc.add_paragraph(style="Caption Text")
            add_rich_text(p, line[1:].strip(), 9, MUTED)
        elif line.startswith("说明："):
            add_note_box(doc, line[3:].strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.keep_together = True
            add_rich_text(p, line)
        i += 1

    doc.core_properties.title = "杭州蔚能海洋科技有限公司宣传稿编辑包"
    doc.core_properties.subject = "杭州良渚新城微信公众号送企业核稿版"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "良渚新城, 企业宣传稿, 蔚能海洋"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

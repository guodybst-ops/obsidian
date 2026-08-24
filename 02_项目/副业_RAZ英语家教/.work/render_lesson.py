"""DOCX renderer for RAZ Level L lesson plans."""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_run(paragraph, text: str, *, bold=False, italic=False, color=None,
             highlight=None, size=None):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    if highlight:
        run.font.highlight_color = highlight
    return run


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def render_lesson(data: dict, out_path: str) -> None:
    doc = Document()
    _set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    # ---- Title block ----
    title = doc.add_heading(data["title_en"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        sub,
        f"Level L  |  Word Count: {data['word_count']}  |  主题：{data['topic']}",
        bold=True,
    )

    # ---- 1. Teaching objectives ----
    doc.add_heading("一、教学目标", level=1)
    doc.add_heading("（一）知识与技能", level=2)
    for item in data["objectives"]["knowledge"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("（二）过程与方法", level=2)
    for item in data["objectives"]["process"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("（三）情感态度与价值观", level=2)
    for item in data["objectives"]["emotions"]:
        doc.add_paragraph(item, style="List Bullet")

    # ---- 2. Key & difficult points ----
    doc.add_heading("二、教学重点与难点", level=1)
    doc.add_paragraph("重点：" + "；".join(data["key_points"]))
    doc.add_paragraph("难点：" + "；".join(data["difficult_points"]))

    # ---- 3. Methods & preparation ----
    doc.add_heading("三、教学方法与准备", level=1)
    doc.add_paragraph("教学方法：" + "，".join(data["methods"]))
    doc.add_paragraph("教学准备：" + "，".join(data["preparation"]))

    # ---- 4. Teaching process ----
    doc.add_heading("四、教学过程", level=1)
    for stage in data["process"]:
        doc.add_heading(f"{stage['name']}（{stage['duration']}）", level=2)
        p1 = doc.add_paragraph()
        _add_run(p1, "教师活动：", bold=True)
        _add_run(p1, stage["teacher"])
        p2 = doc.add_paragraph()
        _add_run(p2, "学生活动：", bold=True)
        _add_run(p2, stage["students"])
        p3 = doc.add_paragraph()
        _add_run(p3, "设计意图：", bold=True)
        _add_run(p3, stage["purpose"])

    # ---- 5. Blackboard ----
    doc.add_heading("五、板书设计", level=1)
    doc.add_paragraph(data["blackboard"])

    # ---- 6. Homework ----
    doc.add_heading("六、作业", level=1)
    for hw in data["homework"]:
        doc.add_paragraph(hw, style="List Number")

    # ---- 7. Vocabulary table ----
    doc.add_heading("七、重点词汇（含英文释义，仅收录六年级课标外生词）", level=1)
    v_table = doc.add_table(rows=1, cols=5)
    v_table.style = "Light Grid Accent 1"
    headers = ["单词", "音标", "词性", "中文释义", "英文释义"]
    for i, h in enumerate(headers):
        cell = v_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_run(p, h, bold=True)
        _set_cell_bg(cell, "DCE6F1")
    for word in data["vocabulary"]:
        row = v_table.add_row().cells
        # Word cell: bold + yellow highlight to mark as 重点
        c0 = row[0]
        c0.text = ""
        p0 = c0.paragraphs[0]
        _add_run(
            p0,
            word["word"],
            bold=True,
            highlight=WD_COLOR_INDEX.YELLOW,
        )
        row[1].text = word["ipa"]
        row[2].text = word["pos"]
        row[3].text = word["meaning_zh"]
        row[4].text = word["meaning_en"]

    # ---- 8. Phrases table ----
    doc.add_heading("八、重点短语", level=1)
    p_table = doc.add_table(rows=1, cols=4)
    p_table.style = "Light Grid Accent 1"
    p_headers = ["短语", "中文释义", "用法说明", "例句"]
    for i, h in enumerate(p_headers):
        cell = p_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_run(p, h, bold=True)
        _set_cell_bg(cell, "DCE6F1")
    for ph in data["phrases"]:
        row = p_table.add_row().cells
        c0 = row[0]
        c0.text = ""
        p0 = c0.paragraphs[0]
        _add_run(
            p0,
            ph["phrase"],
            bold=True,
            highlight=WD_COLOR_INDEX.YELLOW,
        )
        row[1].text = ph["meaning_zh"]
        row[2].text = ph["usage"]
        row[3].text = ph["example"]

    # ---- 9. Grammar & sentence patterns ----
    doc.add_heading("九、重点语法与句式", level=1)
    for gr in data["grammar"]:
        doc.add_heading(gr["point"], level=2)
        p1 = doc.add_paragraph()
        _add_run(p1, "【讲解】", bold=True, color="2F5496")
        _add_run(p1, gr["explanation"])
        p2 = doc.add_paragraph()
        _add_run(p2, "【例句】", bold=True, color="2F5496")
        for ex in gr["examples"]:
            p_ex = doc.add_paragraph(ex, style="Intense Quote")

    doc.save(out_path)
